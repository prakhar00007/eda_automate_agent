"""
export_helpers.py - Report generation and export helper functions
Contains functions for HTML, Word, and CSV export functionality
"""

import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64


def generate_html_report(df, dataset_info):
    """Generate an HTML report of the dataset analysis"""
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>EDA Report - {timestamp}</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    margin: 20px;
                    background-color: #f5f5f5;
                }}
                .container {{
                    background-color: white;
                    padding: 30px;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }}
                h1 {{
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #34495e;
                    margin-top: 30px;
                    border-left: 4px solid #3498db;
                    padding-left: 10px;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 15px 0;
                }}
                th {{
                    background-color: #3498db;
                    color: white;
                    padding: 12px;
                    text-align: left;
                }}
                td {{
                    border: 1px solid #ddd;
                    padding: 10px;
                }}
                tr:nth-child(even) {{
                    background-color: #f9f9f9;
                }}
                .metric {{
                    display: inline-block;
                    margin: 10px 20px 10px 0;
                    padding: 15px;
                    background-color: #ecf0f1;
                    border-radius: 5px;
                }}
                .metric-label {{
                    font-weight: bold;
                    color: #2c3e50;
                }}
                .metric-value {{
                    font-size: 18px;
                    color: #3498db;
                    margin-top: 5px;
                }}
                .timestamp {{
                    text-align: right;
                    color: #7f8c8d;
                    margin-top: 20px;
                    font-size: 12px;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1>Automated EDA Report</h1>
                
                <h2>Dataset Overview</h2>
                <div class="metric">
                    <div class="metric-label">Total Rows</div>
                    <div class="metric-value">{dataset_info['shape'][0]}</div>
                </div>
                <div class="metric">
                    <div class="metric-label">Total Columns</div>
                    <div class="metric-value">{dataset_info['shape'][1]}</div>
                </div>
                <div class="metric">
                    <div class="metric-label">Missing Values</div>
                    <div class="metric-value">{sum(dataset_info['missing_values'].values())}</div>
                </div>
                <div class="metric">
                    <div class="metric-label">Duplicate Rows</div>
                    <div class="metric-value">{dataset_info['duplicates']}</div>
                </div>
                <div class="metric">
                    <div class="metric-label">Memory Usage</div>
                    <div class="metric-value">{dataset_info['memory_usage']}</div>
                </div>

                <h2>Column Information</h2>
                <table>
                    <tr>
                        <th>Column</th>
                        <th>Data Type</th>
                        <th>Non-Null Count</th>
                        <th>Missing Values</th>
                        <th>Missing %</th>
                    </tr>
        """
        
        for idx, (col, dtype) in enumerate(dataset_info['dtypes'].items()):
            missing = dataset_info['missing_values'].get(col, 0)
            missing_pct = dataset_info['missing_percentage'].get(col, 0)
            non_null = dataset_info['shape'][0] - missing
            html_content += f"""
                    <tr>
                        <td>{col}</td>
                        <td>{dtype}</td>
                        <td>{non_null}</td>
                        <td>{missing}</td>
                        <td>{missing_pct:.2f}%</td>
                    </tr>
            """
        
        html_content += """
                </table>

                <h2>Data Sample (First 10 rows)</h2>
        """
        
        html_content += df.head(10).to_html(classes='sample-table', border=0)
        html_content += f"""
                <div class="timestamp">
                    Report generated on: {timestamp}
                </div>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    except Exception as e:
        st.error(f"Error generating HTML report: {str(e)}")
        return None


def generate_word_report(df, dataset_info):
    """Generate a Word document report of the dataset analysis"""
    try:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = Document()
        
        # Title
        title = doc.add_heading('Automated EDA Report', 0)
        title_format = title.paragraph_format
        title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Timestamp
        timestamp_para = doc.add_paragraph(f"Generated on: {timestamp}")
        timestamp_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        timestamp_para.runs[0].font.size = Pt(10)
        timestamp_para.runs[0].font.italic = True
        
        # Dataset Overview Section
        doc.add_heading('Dataset Overview', level=1)
        overview_table = doc.add_table(rows=6, cols=2)
        overview_table.style = 'Light Grid Accent 1'
        
        overview_data = [
            ('Metric', 'Value'),
            ('Total Rows', str(dataset_info['shape'][0])),
            ('Total Columns', str(dataset_info['shape'][1])),
            ('Missing Values', str(sum(dataset_info['missing_values'].values()))),
            ('Duplicate Rows', str(dataset_info['duplicates'])),
            ('Memory Usage', str(dataset_info['memory_usage']))
        ]
        
        for i, (metric, value) in enumerate(overview_data):
            row = overview_table.rows[i]
            row.cells[0].text = metric
            row.cells[1].text = value
            if i == 0:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        # Column Information Section
        doc.add_heading('Column Information', level=1)
        col_table = doc.add_table(rows=len(dataset_info['dtypes']) + 1, cols=5)
        col_table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = col_table.rows[0].cells
        headers = ['Column', 'Data Type', 'Non-Null Count', 'Missing Values', 'Missing %']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for idx, (col, dtype) in enumerate(dataset_info['dtypes'].items(), 1):
            row = col_table.rows[idx]
            missing = dataset_info['missing_values'].get(col, 0)
            missing_pct = dataset_info['missing_percentage'].get(col, 0)
            non_null = dataset_info['shape'][0] - missing
            
            row.cells[0].text = col
            row.cells[1].text = str(dtype)
            row.cells[2].text = str(non_null)
            row.cells[3].text = str(missing)
            row.cells[4].text = f"{missing_pct:.2f}%"
        
        # Data Sample Section
        doc.add_heading('Data Sample (First 10 Rows)', level=1)
        sample_df = df.head(10)
        sample_table = doc.add_table(rows=len(sample_df) + 1, cols=len(sample_df.columns))
        sample_table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = sample_table.rows[0].cells
        for i, col in enumerate(sample_df.columns):
            header_cells[i].text = str(col)
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for row_idx, row in sample_df.iterrows():
            cells = sample_table.rows[row_idx + 1].cells
            for col_idx, value in enumerate(row):
                cells[col_idx].text = str(value)
        
        return doc
    
    except Exception as e:
        st.error(f"Error generating Word report: {str(e)}")
        return None


def download_csv_report(df, dataset_info):
    """Generate a CSV summary of the dataset"""
    try:
        csv_data = df.to_csv(index=False)
        return csv_data
    except Exception as e:
        st.error(f"Error generating CSV report: {str(e)}")
        return None


def create_download_button(file_content, file_name, file_type, label):
    """Create a download button for various file types"""
    try:
        if file_type == "html":
            b64 = base64.b64encode(file_content.encode()).decode()
            href = f'<a href="data:text/html;base64,{b64}" download="{file_name}">Download {label}</a>'
            return href
        
        elif file_type == "csv":
            b64 = base64.b64encode(file_content.encode()).decode()
            href = f'<a href="data:text/csv;base64,{b64}" download="{file_name}">Download {label}</a>'
            return href
        
        elif file_type == "docx":
            # For DOCX, we need to save to bytes
            return None  # Return None, let app.py handle the byte conversion
        
        return None
    
    except Exception as e:
        st.error(f"Error creating download button: {str(e)}")
        return None
